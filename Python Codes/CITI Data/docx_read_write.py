from docx import Document

def read_from_docx(path):
    doc = Document(path)
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text
    return all_text

extracted_text = read_from_docx("Python Codes\Intership\sample.docx")
print(extracted_text)

#creating new document
doc = Document()
doc.add_heading('Document Title', 0)
doc.add_paragraph("This is paragraph added through python")
doc.save("Python Codes\Intership\output.docx")